#!/usr/bin/env python3
"""Generate the AirStack->Starling milestones Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---------- styles ----------
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

def code_block(text):
    for line in text.strip('\n').split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x6E)
    doc.add_paragraph()

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

# ---------- title ----------
t = doc.add_heading('AirStack + Starling Max 2 Live Flight', level=0)
sub = doc.add_paragraph()
r = sub.add_run('Milestone Plan, Progress Log, and Replication Runbook')
r.italic = True
r.font.size = Pt(13)
para('Last updated: 20 July 2026  |  Branch: daniel/diffaero_ground_control  |  Working copy: ~/AirStack-diffaero')

# ---------- 1 objective ----------
doc.add_heading('1. Objective', level=1)
para('Replicate CMU AirLab’s ground-controller workflow on our own hardware: fly a ModalAI '
     'Starling Max 2 live in the real world, commanded by AirStack running on a ground laptop, '
     'using our OptiTrack + Motive motion-capture system as the drone’s indoor position source. '
     'The CBF collision-avoidance scenarios and multi-drone swarm demos that CMU also built are '
     'explicitly OUT of scope for now — we only need single-drone takeoff, hover, and landing '
     'under mocap. The stack we use is CMU’s own flight-tested branch '
     '(daniel/diffaero_ground_control), which already contains every piece of this pipeline: '
     'the NatNet mocap driver (natnet_ros2), the mocap-to-PX4 bridge (mocap_bridge), the '
     'uXRCE-DDS flight interface (px4_interface + MicroXRCEAgent), and the swarm commander that '
     'provides takeoff / hold / land services and a software geofence.')

# ---------- 2 purpose ----------
doc.add_heading('2. What the Milestones Are For', level=1)
para('The milestones are a debugging strategy applied before anything breaks. The flight-day system is a '
     'long chain: Motive → natnet_ros2 → mocap_bridge → XRCE agent → WiFi → PX4 EKF2 '
     '→ px4_interface → swarm commander → motors. If it were assembled in one go and failed, '
     'the fault could be in any of eight links. Each milestone instead adds exactly ONE new link and proves '
     'it works before the next one stacks on top, so a failure always points at the link just added.')
bullet('Sim rehearsal (M1) proves the software and the operator — the simulated drones run the real '
       'PX4 firmware (SITL), so the arm/OFFBOARD/setpoint negotiation is genuinely exercised, and the '
       'operator learns the cockpit for free.', 'M1: ')
bullet('Props-off stages (M3, M4) prove the comms and estimation links with the drone physically unable '
       'to hurt anyone.', 'M3–M4: ')
bullet('Hand-carry (M5) is the full-system dress rehearsal minus the motors: every flight-day node '
       'running, the stack’s belief checked against reality.', 'M5: ')
bullet('First flight (M6) then introduces exactly one untested variable: thrust.', 'M6: ')
para('Exit criteria are evidence, not vibes — each is a fact observed on screen. Milestones are also '
     're-entry points: a lab day cut short at M3 resumes at M4 with nothing re-derived.')

# ---------- 3 status ----------
doc.add_heading('3. Status at a Glance', level=1)
table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Milestone'
hdr[1].text = 'Goal'
hdr[2].text = 'Status'
rows = [
    ('M1 — Sim rehearsal', '3 SITL drones fly under the ground controller in Isaac Sim; teleop, geofence breach + recovery exercised', 'COMPLETE (20 Jul 2026)'),
    ('M2 — Ground station hardware prep', 'Host networking, Motive config, time sync, port checks', 'Partial — see §6'),
    ('M3 — Drone comms (props off)', 'Real PX4 topics streaming to the laptop over WiFi (uXRCE-DDS)', 'Not started (needs drone)'),
    ('M4 — Mocap → EKF2 (props off)', 'OptiTrack pose fused by the drone’s EKF2; frames verified', 'Not started (needs drone + mocap room)'),
    ('M5 — Hand-carry preflight', 'RViz marker tracks the hand-carried drone through the volume', 'Not started'),
    ('M6 — First flight', 'Stable mocap-fused hover + landing under AirStack command', 'Not started'),
]
for i, (a, b, c) in enumerate(rows, start=1):
    cells = table.rows[i].cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c
doc.add_paragraph()

# ---------- 4 M1 log ----------
doc.add_heading('4. Milestone 1 — What Happened (Completed 20 July 2026)', level=1)
doc.add_heading('4.1 Setup work that had to happen first', level=2)
bullet('Cloned CMU’s branch to ~/AirStack-diffaero (git clone -b daniel/diffaero_ground_control ... plus git submodule update --init).')
bullet('Copied two gitignored config files a fresh clone never gets: simulation/isaac-sim/docker/omni_pass.env '
       '(Omniverse credentials) and user.config.json (Isaac Kit config). Symptom when missing: “env file ... not found” '
       'on up, and a Docker error about mounting a directory onto a file (Docker auto-creates a missing mount source as an '
       'empty directory — delete the bogus directory, copy the real file).')
bullet('Rebuilt the robot Docker image (./airstack.sh image-build robot-desktop). REQUIRED on this branch: it bakes in '
       'MicroXRCEAgent (the uXRCE-DDS agent for real drones) and pins ROS_DOMAIN_ID=1 in the container .bashrc. '
       'A plain up with the stock v0.18.0 image is missing both.')
bullet('Re-applied the ZED camera-info race fix to the clone’s PegasusSimulator submodule '
       '(spawn_zed_camera.py line 224: trigger the info helper from right_create_rp.outputs:execOut instead of the '
       'playback tick). Any fresh checkout needs this until CMU’s fix/camera-init PR merges.')
bullet('Built the ROS 2 workspace inside the robot container with bws — 59 packages in ~4 min '
       '(px4_msgs dominates). Warnings are normal; only “failed” matters.')

doc.add_heading('4.2 The bring-up that flew', level=2)
para('Five terminals, each with one job:')
bullet('Isaac container — spawn script svg_multi_drone_single_domain.py spawned 3 drones with PX4 SITL '
       'each, on ROS domain 1, playing immediately, headless. Success line: “Spawning 3 drone(s) on ROS domain 1”, '
       'then “Ready for takeoff!” ×3 from the PX4 commanders.', 'Terminal 1 (sim): ')
bullet('Robot container — bws, sws, then launch_sim_interfaces.sh 3 started one MAVROS instance per '
       'drone. Success: “Got HEARTBEAT, connected. FCU: PX4 Autopilot” and mission/rallypoint/geofence '
       'handshakes for all three drones; /drone_1/interface/mavros/state showed connected: true.', 'Terminal 2 (interfaces): ')
bullet('Robot container — ros2 launch svg_ground_control ground_control.launch.py. The SwarmCommander '
       'came up: scenario=hover, drone_1/2 autonomous, drone_3 teleop, CBF r=0.55 m vmax=1.2 m/s, fence '
       '[-4,-2,0]..[4,2,3]. (mocap_bridge also starts and idles harmlessly in sim.)', 'Terminal 3 (commander): ')
bullet('Robot container — rviz2 with svg_drones.rviz shows all drones as spheres on /svg/viz/markers.', 'Terminal 4 (RViz): ')
bullet('Robot container — service calls: /swarm_commander/takeoff then start — the drones armed, '
       'climbed, and moved. Milestone complete.', 'Terminal 5 (cockpit): ')

doc.add_heading('4.3 Lessons learned (the traps)', level=2)
bullet('Host shell vs container shell: ros2 commands only work INSIDE the robot container '
       '(./airstack.sh connect robot --command=bash; prompt changes to root@). A host shell is on the wrong '
       'DDS domain — topic lists come back near-empty and service calls hang at “waiting for service”. '
       'Rule: prompt jeremychia@ = laptop (wrong), root@ = container (right).')
bullet('Paste one line at a time. connect opens an interactive shell; everything pasted after it in the same block '
       'gets typed into whichever shell is ready, causing “command not found” chaos and garbage stuck in the '
       'input buffer (Ctrl+C clears it).')
bullet('bws/sws are AirStack bash functions defined in robot/docker/.bashrc — bws = colcon build '
       '--symlink-install (with a build lock), sws = source install/local_setup.bash. Code lives on the host '
       '(bind mount), compilation happens in the container, artifacts persist on the host.')
bullet('PX4 [timesync] “time jump detected” warnings in SITL are benign at startup; continuous spam means '
       'the sim is below real-time (heavy GPU load). Not a blocker; reduce load (fewer drones) if drones look wobbly.')
bullet('Container names in this clone are prefixed airstack-diffaero-* (folder name); '
       './airstack.sh connect robot resolves the right one regardless.')
bullet('PlotJuggler or any apt-installed tool vanishes on airstack down/up (fresh container from image) but '
       'survives airstack restart. Reinstall per session or bake into the image.')

doc.add_heading('4.4 Additional tests performed, incidents, and fixes (later the same day)', level=2)
bullet('Commander state machine learned the hard way: teleop input is IGNORED until the scenario is started. '
       'Flow: IDLE --takeoff--> HOLDING (holds takeoff position, ignores nominal inputs) --start--> ACTIVE '
       '(auto drones fly the scenario, teleop drones obey /svg/<name>/teleop_command) --hold--> HOLDING (panic '
       'freeze). Symptom of forgetting start: teleop topic shows changing values but the drone does not move.', 'Teleop: ')
bullet('keyboard_teleop reads keypresses from ITS OWN terminal — click that terminal for focus. '
       'w/s = ±x, a/d = ±y, r/f = up/down, space = zero command, +/- = speed, q = quit.', 'Teleop focus: ')
bullet('drone_3 was teleop-flown through the fence wall (y<min). The commander latched the breach exactly as designed: '
       'ALL drones froze in place (orange in RViz), scenario stopped, start refused, fence box drawn red. This is the '
       'same behavior CMU validated on real hardware. Recovery used: land (fence-exempt) → reset_fence → takeoff → start. '
       'Key hardware note: the fence FREEZES, it does not cut motors — the RC kill switch is the only true cutoff.', 'Geofence breach (accidental, valuable): ')
bullet('During breach recovery, drone_3’s land command FAILED (still outside the fence) — the first failed service '
       'report ever, which took a never-exercised error path in swarm_commander.py report() (~line 588): '
       '“level = logger.info if ok else logger.error; level(msg)”. rclpy caches log severity per call-site, so one '
       'shared line logging at two severities raises ValueError (“Logger severity cannot be changed between calls”) — '
       'killing the commander process WITH AN AIRCRAFT STILL AIRBORNE. FIXED locally: split into separate '
       'if/else info()/error() calls. MUST BE REPORTED UPSTREAM to CMU. Hardware lesson: the ground commander is a '
       'single point of failure — PX4-side offboard-loss failsafe (COM_OBL_RC_ACT) and the RC kill switch are '
       'non-negotiable.', 'Commander crash (real CMU bug, found + fixed): ')
bullet('After ~20 min of continuous SITL hover, PX4 reported “Preflight Fail: Battery unhealthy” — simulated batteries '
       'drain and block re-arming. Reset: Ctrl+C the Isaac spawn script and re-run it (fresh drones, full batteries); '
       'MAVROS reconnects on its own.', 'SITL battery drain: ')

# ---------- 5 replication ----------
doc.add_heading('5. How to Replicate Milestone 1 (Re-run Runbook)', level=1)

doc.add_heading('5.0 One-time setup from scratch (only for a NEW machine or fresh checkout)', level=2)
para('The code lives on CMU’s repo github.com/castacks/AirStack. The branch that contains the entire '
     'ground-controller + mocap + real-drone pipeline is daniel/diffaero_ground_control '
     '(the base SVG branch yikuan/SVG_ground_control also exists, but daniel/... is newer and more complete — '
     'it is what we use). Note: main and develop do NOT contain any of this.')
code_block("""# 1. Clone the branch (submodules pulled separately on purpose --
#    a blind --recurse-submodules fails on private submodules other branches reference)
git clone -b daniel/diffaero_ground_control https://github.com/castacks/AirStack.git ~/AirStack-diffaero
cd ~/AirStack-diffaero && git submodule update --init

# 2. Copy the two gitignored config files from an existing checkout (or recreate from
#    the *_TEMPLATE files next to them: omni_pass_TEMPLATE.env / user_TEMPLATE.config.json)
cp <existing_checkout>/simulation/isaac-sim/docker/omni_pass.env      simulation/isaac-sim/docker/
cp <existing_checkout>/simulation/isaac-sim/docker/user.config.json   simulation/isaac-sim/docker/

# 3. Rebuild the robot image (REQUIRED on this branch: bakes in MicroXRCEAgent, pins domain 1)
./airstack.sh image-build robot-desktop

# 4. Re-apply the camera-race fix (until CMU's fix/camera-init PR is merged):
#    in simulation/isaac-sim/extensions/PegasusSimulator/extensions/pegasus.simulator/
#       pegasus/simulator/ogn/api/spawn_zed_camera.py  (line ~224) change:
#    (f"{nodes['playback']}.outputs:tick",           f"{nodes['info_helper']}.inputs:execIn"),
#    to:
#    (f"{right_nodes['create_rp']}.outputs:execOut", f"{nodes['info_helper']}.inputs:execIn"),

# 5. Fix the swarm_commander logging crash (until reported/merged upstream):
#    in robot/ros_ws/src/svg_ground_control/svg_ground_control/swarm_commander.py,
#    report() (~line 588), replace the two lines
#        level = self.get_logger().info if ok else self.get_logger().error
#        level(f'{name}: {label} -> success={ok}')
#    with separate call-sites:
#        if ok:
#            self.get_logger().info(f'{name}: {label} -> success={ok}')
#        else:
#            self.get_logger().error(f'{name}: {label} -> success={ok}')

# 6. Sanity-check .env at the repo root:
#    COMPOSE_PROFILES="desktop,isaac-sim"   AUTOLAUNCH="false"   NUM_ROBOTS="1" """)

doc.add_heading('5.1 Re-run from a fully stopped system (setup above already done)', level=2)
para('One-time setup (clone, env files, image build, camera fix, first bws) is DONE on this laptop and persists on disk. '
     'A fresh re-run from a fully stopped system is only the following. Run each numbered block in its own '
     'terminal; inside-container blocks are marked. Paste one line at a time.', bold=False)

doc.add_heading('Terminal 1 — stack up (host)', level=2)
code_block("""cd ~/AirStack-diffaero
./airstack.sh up
./airstack.sh status        # robot-desktop-1, gcs, isaac-sim all Up""")

doc.add_heading('Terminal 1 again — Isaac spawn (inside isaac-sim container)', level=2)
code_block("""./airstack.sh connect isaac-sim --command=bash
# then inside (this multi-line block is ONE command, safe to paste whole):
NUM_ROBOTS=3 SVG_DOMAIN_ID=1 PLAY_SIM_ON_START=true ISAAC_SIM_HEADLESS=true \\
PYTHONPATH="$ISAAC_SIM_PYTHONPATH" \\
/isaac-sim/python.sh /isaac-sim/AirStack/simulation/isaac-sim/launch_scripts/svg_multi_drone_single_domain.py \\
  --ext-folder ~/.local/share/ov/data/documents/Kit/shared/exts""")
para('Wait for: “Spawning 3 drone(s) on ROS domain 1” and “Ready for takeoff!” ×3. Leave running. '
     '(Use NUM_ROBOTS=1 for a lighter single-drone rehearsal — also update the commander config to match.)')

doc.add_heading('Terminal 2 — interfaces (inside robot container)', level=2)
code_block("""cd ~/AirStack-diffaero
./airstack.sh connect robot --command=bash
# inside:
cd ~/AirStack/robot/ros_ws
bws        # only needed if code changed since last build; else skip
sws
./src/svg_ground_control/scripts/launch_sim_interfaces.sh 3""")
para('Leave running. Verify from any other container shell: '
     'ros2 topic echo /drone_1/interface/mavros/state --once → connected: true.')

doc.add_heading('Terminal 3 — ground controller (inside robot container)', level=2)
code_block("""cd ~/AirStack-diffaero
./airstack.sh connect robot --command=bash
ros2 launch svg_ground_control ground_control.launch.py""")

doc.add_heading('Terminal 4 — RViz (inside robot container)', level=2)
code_block("""cd ~/AirStack-diffaero
./airstack.sh connect robot --command=bash
rviz2 -d $(ros2 pkg prefix svg_ground_control)/share/svg_ground_control/config/svg_drones.rviz""")

doc.add_heading('Terminal 5 — cockpit (inside robot container)', level=2)
code_block("""cd ~/AirStack-diffaero
./airstack.sh connect robot --command=bash
ros2 service call /swarm_commander/takeoff std_srvs/srv/Trigger
ros2 service call /swarm_commander/start   std_srvs/srv/Trigger
ros2 service call /swarm_commander/hold    std_srvs/srv/Trigger   # panic freeze
ros2 service call /swarm_commander/land    std_srvs/srv/Trigger
# optional teleop for drone_3 (separate shell; requires 'start' to have been called,
# and keypresses go to the teleop terminal itself — click it for focus):
# ros2 run svg_ground_control keyboard_teleop --ros-args -p drone:=drone_3
# geofence breach recovery (fence box red, all drones frozen orange):
# ros2 service call /swarm_commander/land        std_srvs/srv/Trigger
# ros2 service call /swarm_commander/reset_fence std_srvs/srv/Trigger
# then takeoff + start again""")
para('State machine (same on hardware): IDLE —takeoff→ HOLDING —start→ ACTIVE —hold→ HOLDING; land from any '
     'airborne state. HOLDING ignores nominal inputs by design (that is the freeze).')

doc.add_heading('Shutdown', level=2)
code_block("""cd ~/AirStack-diffaero
./airstack.sh down          # or: ./airstack.sh restart <container> to reset one piece""")

# ---------- 6 upcoming milestones ----------
doc.add_heading('6. Upcoming Milestones (To Do)', level=1)
para('Full command detail for every step lives in the branch’s own maintained guide: '
     'robot/ros_ws/src/svg_ground_control/experiment.md (Parts B and D). The steps below follow it, '
     'plus our lab-specific additions. Substitute <LAPTOP_IP> and <MOTIVE_IP>.')

doc.add_heading('M2 — Ground station hardware prep (desk, no drone)', level=2)
para('Already done: QGroundControl AppImage on the laptop; standalone PX4 SITL build (optional rehearsal tool); '
     'PlotJuggler workflow known (reinstall per session or bake into image).', bold=False)
bullet('Robot container onto the LAN: edit robot/docker/docker-compose.yaml (robot-desktop service) — comment out '
       'networks: and ports:, add network_mode: host. Then ./airstack.sh down && ./airstack.sh up robot-desktop. '
       'Required because NatNet and uXRCE-DDS traffic cannot reach a Docker bridge network.')
bullet('Motive PC: rigid body named exactly drone_1 (asymmetric marker layout); Streaming pane: Up Axis = Z, '
       'Broadcast Frame ON, Local Interface = <MOTIVE_IP>.')
bullet('Time sync: chrony/NTP across laptop + Motive PC (VOXL joins in M3).')
bullet('Port check (lesson from our previous mocap outage): ss -ulpn | grep -E ‘1510|1511’ — must be clear '
       'of orphan processes before every mocap session.')

doc.add_heading('M3 — Drone comms, props off', level=2)
bullet('Drone on LAN via adb: check ip addr show wlan0; use udhcpc / voxl-wifi station for DHCP; '
       'prefer a router DHCP reservation per drone.')
bullet('Push and run CMU’s one-shot comms script on the VOXL: adb push .../voxl_setup_real_drone.sh /usr/bin/ '
       'then voxl_setup_real_drone.sh drone_1 <LAPTOP_IP> 1 8888. It points PX4’s DDS client at the laptop, '
       'namespaces topics /drone_1/fmu/..., pins domain 1, disables the onboard voxl-microdds-agent, restarts voxl-px4. '
       'Verify on VOXL: px4-microdds_client status → connected, Agent IP = <LAPTOP_IP>.')
bullet('Ground agent (robot container, leave running): MicroXRCEAgent udp4 -p 8888 -v4 — watch for '
       '“session established”.')
bullet('Verify topics arrive: ros2 topic list | grep drone_1/fmu ; echo vehicle_status WITH '
       '--qos-reliability best_effort (PX4 topics are best-effort — without the flag they look dead).')

doc.add_heading('M4 — Mocap into EKF2, props off', level=2)
bullet('Build + launch the OptiTrack driver (robot container): bws --packages-select natnet_ros2 (first build '
       'downloads the NatNet SDK — needs internet), then ros2 launch natnet_ros2 natnet_ros2.launch.py '
       'serverIP:=<MOTIVE_IP> clientIP:=<LAPTOP_IP>. NOTE: launch defaults are CMU’s lab IPs — always override.')
bullet('Unit-test the stream: ros2 topic hz /drone_1/pose (≈ Motive rate); echo once; hand-carry — smooth, no jumps.')
bullet('PX4 EKF2 params, once per drone (QGC or px4-param): EKF2_EV_CTRL=11, EKF2_HGT_REF=3, EKF2_GPS_CTRL=0, '
       'EKF2_EV_DELAY≈50. Without these PX4 ignores external vision, produces no estimate, and REFUSES TO ARM '
       'indoors (“fuse failure”) — mocap fusion is mandatory, not optional.')
bullet('Audit for a second vision source: voxl-inspect-services — if voxl-vision-hub feeds VIO into PX4, disable '
       'that feed. One external-vision source at a time.')
bullet('Start commander + mocap bridge (no takeoff): ros2 launch svg_ground_control ground_control.launch.py '
       'config:=.../swarm_real.yaml use_mocap:=true. Verify /drone_1/fmu/in/vehicle_visual_odometry streams and '
       '/drone_1/fmu/out/vehicle_odometry appears (best_effort QoS) — that is EKF2 fusing.')
bullet('FRAME HAND-CHECK (before every first flight): carry 1 m toward North → position[0] increases; East → '
       'position[1] increases; lift → position[2] DEcreases (NED). Mirrored? Set px4_vio_frame: "modalai_flip" in '
       'swarm_real.yaml and re-check. A wrong frame flies the drone into a wall.')

doc.add_heading('M5 — Hand-carry preflight (nothing armed)', level=2)
bullet('All flight-day nodes running: agent, natnet, real_interfaces (ros2 launch svg_ground_control '
       'real_interfaces.launch.py drones:=drone_1), commander (idle — do NOT call takeoff), RViz.')
bullet('Carry the drone around the volume: its red sphere in RViz must track reality everywhere.')
bullet('Record a ground-truth bag: ros2 bag record /drone_1/pose /drone_1/odometry_conversion/odometry.')

doc.add_heading('M6 — First flight', level=2)
bullet('Config (swarm_real.yaml): drone_names: ["drone_1"], drone_modes: "real", scenario hover, fence_min/max '
       'fitted INSIDE the net volume, speed cap ≤ 1.0 m/s.')
bullet('PX4 safety: RC kill switch mapped and tested; offboard-loss failsafe (COM_OBL_RC_ACT); low-battery action. '
       'The software geofence only FREEZES — the RC kill switch is the only true motor cutoff.')
bullet('Preflight: mocap hz, odometry tracks reality, thumb on kill switch.')
bullet('Fly: takeoff → hover → land via the same three service calls as sim. Recover a fence latch with '
       '/swarm_commander/reset_fence.')
bullet('Post-flight: PlotJuggler diff of /drone_1/pose vs /drone_1/fmu/out/vehicle_odometry from the bag.')

# ---------- 7 gotchas ----------
doc.add_heading('7. Quick Troubleshooting Reference', level=1)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Symptom'
tbl.rows[0].cells[1].text = 'Cause / Fix'
trows = [
    ('ros2: command not found, or topic list nearly empty, or service call hangs “waiting”',
     'You are in a HOST shell. ./airstack.sh connect robot --command=bash first (prompt must say root@).'),
    ('“env file omni_pass.env not found” on up',
     'Gitignored file missing in a fresh clone — copy from another checkout.'),
    ('Docker: “mounting ... user.config.json ... not a directory”',
     'A failed up auto-created the missing file as a directory. rmdir it, copy the real file, up again.'),
    ('bws: command not found',
     'Wrong container (Isaac has no bws) or host shell. bws exists only in the robot container.'),
    ('/fmu/* topics look dead',
     'Best-effort QoS — add --qos-reliability best_effort to echo/hz.'),
    ('PX4 refuses to arm indoors (“fuse failure”)',
     'EKF2 has no position source — mocap feed or EKF2_EV_* params missing (M4).'),
    ('Mocap topic silent / 0 Hz',
     'Motive not streaming, wrong serverIP, rigid body not named drone_N, or an orphan process on UDP 1510/1511.'),
    ('Continuous [timesync] warnings in sim',
     'Sim below real-time (GPU load). Reduce NUM_ROBOTS / close viewers. Sim-only issue.'),
    ('Teleop topic publishes but drone does not move',
     'Commander is in HOLDING — call /swarm_commander/start first; also click the teleop terminal for keyboard focus.'),
    ('GEOFENCE BREACH — all drones frozen, fence red',
     'By design. land (fence-exempt) → reset_fence → takeoff → start. Fence freezes only; RC kill is the real cutoff.'),
    ('swarm_commander dies: “Logger severity cannot be changed between calls”',
     'CMU bug in report() (~line 588), triggered by the first FAILED service report. Fixed locally (split info/error call-sites); rebuild svg_ground_control. Report upstream.'),
    ('PX4 “Preflight Fail: Battery unhealthy” in sim; refuses to arm',
     'SITL virtual battery drained after long hover. Ctrl+C the Isaac spawn script and re-run it.'),
]
for a, b in trows:
    row = tbl.add_row()
    row.cells[0].text = a
    row.cells[1].text = b
doc.add_paragraph()

# ---------- 8 references ----------
doc.add_heading('8. References', level=1)
bullet('Canonical command reference (CMU, maintained): ~/AirStack-diffaero/robot/ros_ws/src/svg_ground_control/experiment.md '
       '— Part A (sim), Part B (real drone bring-up), Part C (tasks), Part D (first flight), Geofence, Troubleshooting.')
bullet('Package overview: ~/AirStack-diffaero/robot/ros_ws/src/svg_ground_control/README.md')
bullet('Mocap driver: robot/ros_ws/src/perception/natnet_ros2 (NatNet → /drone_N/pose, ports 1510/1511)')
bullet('Mocap→PX4 bridge: robot/ros_ws/src/svg_ground_control/svg_ground_control/mocap_bridge.py '
       '(timestamp=0 restamping trick, enu_to_ned vs modalai_flip frames)')
bullet('VOXL comms script: robot/ros_ws/src/svg_ground_control/scripts/voxl_setup_real_drone.sh')

doc.save('/home/jeremychia/AirStack_Starling_Milestones.docx')
print('saved')
